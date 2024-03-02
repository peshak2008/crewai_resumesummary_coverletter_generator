import os
import json
from dotenv import load_dotenv
from docx import Document
from crewai import Agent, Task, Crew, Process
from langchain_openai import ChatOpenAI
from langchain.tools import tool
from crewai.tasks.task_output import TaskOutput
from crewai_tools.tools import FileReadTool
from utils.config import candidate_name, resume_name, resume_template,default_job_title,default_skills_summary,default_location,llm_model,create_pdf,resume_sections_to_update

        
class ReadDescription():
        
    @tool("Read Description Tool")
    def read_description(job_index, file_path):
        """Read a specific job description from an array in a file."""
        print(f"Reading job description for job index {job_index}")
        with open(file_path, 'r') as file:
            content = json.load(file)
            return content[job_index]
            

            
class JobSearchCrewManager:
    def __init__(self, job_index):
        load_dotenv()
        self.llm = ChatOpenAI(model=llm_model)
        self.default_resume_folder = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'default_resume')
        self.resume_md_name = os.path.splitext(resume_name)[0] + '.md'
        self.resume_read_tool = FileReadTool(file_path=os.path.join(self.default_resume_folder, self.resume_md_name))
        self.job_details = {'job_title': "unknown",'company_name': "unknown", "job_location": default_location} #initialize job details
        self.job_index = job_index  # Store the job indices
        print(f"Job indices: {job_index}")
       
        self.read_description_tool = ReadDescription.read_description

        self.output_file = "output.json"


    def create_custom_resume_docx(self, template_path, new_file_path, job_title, summary, job_location):
        doc = Document(template_path)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if '{TITLE}' in run.text and resume_sections_to_update.get('title', False):
                    run.text = run.text.replace('{TITLE}', job_title)
                elif '{SUMMARY}' in run.text and resume_sections_to_update.get('summary', False):
                    run.text = run.text.replace('{SUMMARY}', summary)
                elif '{LOCATION}' in run.text and resume_sections_to_update.get('location', False):
                    run.text = run.text.replace('{LOCATION}', job_location)
        doc.save(new_file_path)


    def convert_docx_to_pdf(self, docx_path, pdf_path):
        import subprocess
        cmd = ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), docx_path]
        subprocess.run(cmd, check=True)

    def job_callback_function(self,output: TaskOutput):
        global job_details  # Indicate that we're using the global variable
        content = output.raw_output

        job_title = None
        company_name = None
        job_location = None
        role_seniority = "unknown"  # Default value
        role_discipline = "unknown"  # Default value
        company_industry = "unknown"  # Default value
        salary_range = "unknown"  # Default value

        lines = [line.strip() for line in content.splitlines()] # Split the content into lines and remove leading/trailing whitespace

        for line in lines:
            if line.startswith("Job Title:"):
                job_title = line.replace("Job Title:", "").strip()
            elif line.startswith("Company Name:"):
                company_name = line.replace("Company Name:", "").strip()
            elif line.startswith("Job Location:"):
                job_location = line.replace("Job Location:", "").strip()
            elif line.startswith("Role Seniority:"):
                role_seniority = line.replace("Role Seniority:", "").strip()
            elif line.startswith("Role Discipline:"):
                role_discipline = line.replace("Role Discipline:", "").strip()
            elif line.startswith("Company Industry:"):
                company_industry = line.replace("Company Industry:", "").strip()
            elif line.startswith("Salary Range:"):
                salary_range = line.replace("Salary Range:", "").strip()


        # If job_title or company_name  or job_location were not found, set them as 'unknown'
        job_title = job_title if job_title is not None else 'unknown'
        company_name = company_name if company_name is not None else 'unknown'
        job_location = job_location if job_location is not None else default_location
        role_seniority = role_seniority if role_seniority is not None else 'unknown'
        role_discipline = role_discipline if role_discipline is not None else 'unknown'
        company_industry = company_industry if company_industry is not None else 'unknown'
        salary_range = salary_range if salary_range is not None else 'unknown'



        # Store the extracted values in the global dictionary
        self.job_details['job_title'] = job_title
        self.job_details['company_name'] = company_name
        self.job_details['job_location'] = job_location
        self.job_details['role_seniority'] = role_seniority
        self.job_details['role_discipline'] = role_discipline
        self.job_details['company_industry'] = company_industry
        self.job_details['salary_range'] = salary_range

        # Print extracted information
        print(f"Extracted Job Title: {job_title}")
        print(f"Extracted Company Name: {company_name}")
        print(f"Extracted Job Location: {job_location}")

        company_name = str(company_name)  # Convert to string if it's not
        job_title = str(job_title)        # Convert to string if it's not
        job_location = str(job_location)  # Convert to string if it's not

        # Create a folder with the company name and a nested folder with the job title
        nested_folder_path = os.path.join(os.getcwd(), "jobs",company_name, job_title)
        os.makedirs(nested_folder_path, exist_ok=True)  # Avoids error if the folder already exists

        # Define the file path for the job description summary inside the nested folder
        file_path = os.path.join(nested_folder_path, "job_description_summary.txt")

        # Write the content to the file
        try:
            with open(file_path, "w") as file:
                file.write(content)
            print(f"Job description saved to {file_path}")
        except IOError as e:
            print(f"Failed to write file {file_path}: {e}")

    def candidate_callback_function(self,output: TaskOutput):
        global job_details  # Ensure we're using the global dictionary

        # Extract the raw output from the task
        content = output.raw_output

    # Retrieve company name and job title from the global dictionary, replace with "unknown" if missing
        company_name = self.job_details.get('company_name', 'unknown')
        job_title = self.job_details.get('job_title', 'unknown')

        # Construct the nested folder path using company name and job title
        nested_folder_path = os.path.join(os.getcwd(),"jobs", company_name, job_title)

        # Check if the nested folder path exists
        if not os.path.exists(nested_folder_path):
            print(f"Error: The folder path {nested_folder_path} does not exist.")
            return

        # Define the file path for the candidate skills summary inside the nested folder
        file_path = os.path.join(nested_folder_path, "candidate_skills_summary.txt")

        # Write the content to the file
        try:
            with open(file_path, "w") as file:
                file.write(content)
            print(f"Candidate skills summary saved to {file_path}")
        except IOError as e:
            print(f"Failed to write file {file_path}: {e}")

    def resume_callback_function(self,output: TaskOutput):
        global job_details  # Ensure we're using the global dictionary

        print(self.job_details.get('company_name','unknown'))
        print(self.job_details.get('job_title','unknown'))
        print(self.job_details.get('job_location','unknown'))

        # Extract the raw output from the task
        content = output.raw_output

        # Retrieve company name and job title from the global dictionary, replace with "unknown" if missing
        company_name = self.job_details.get('company_name', 'unknown')
        job_title = self.job_details.get('job_title', 'unknown')
        job_location = self.job_details.get('job_location', 'unknown')

        # Construct the nested folder path using company name and job title
        nested_folder_path = os.path.join(os.getcwd(), "jobs", company_name, job_title)

        # Check if the nested folder path exists
        if not os.path.exists(nested_folder_path):
            print(f"Error: The folder path {nested_folder_path} does not exist.")
            return

        # Define the file path for the resume summary inside the nested folder
        file_path = os.path.join(nested_folder_path, "resume_summary.txt")

        # Write the content to the file
        try:
            with open(file_path, "w") as file:
                file.write(content)
            print(f"Resume summary saved to {file_path}")
        except IOError as e:
            print(f"Failed to write file {file_path}: {e}")

        # Define the file path for the custom resume inside the nested folder
        custom_resume_path = os.path.join(nested_folder_path, f"Resume - {candidate_name} - {company_name} - {job_title}.docx")

        # Define the path to the default_resume folder relative to this script
        template_path = os.path.join(self.default_resume_folder, resume_template)

        # Clean the output by removing unexpected characters
        clean_output = content.replace("```", "")

        lines = clean_output.splitlines()

        # Remove unnecessary blank lines
        lines = [line for line in lines if line.strip()]

        # Initialize variables
        extracted_job_title = default_job_title
        extracted_summary = default_skills_summary

        # Find indices for "Job Title:" and "Summary:"
        job_title_index = None
        summary_index = None

        for i, line in enumerate(lines):
            if line.startswith("Job Title:"):
                job_title_index = i
            elif "Summary:" in line:
                summary_index = i
                break  # Stop searching once "Summary:" is found

        # Extract job title if found
        if job_title_index is not None and summary_index is not None:
            extracted_job_title = lines[job_title_index].replace("Job Title:", "").strip()

        # Extract summary if found
        if summary_index is not None:
            if len(lines[summary_index].split(":")) > 1:
                # Summary is on the same line as "Summary:"
                extracted_summary = lines[summary_index].split(":", 1)[1].strip() + "\n" + "\n".join(lines[summary_index + 1:])
            else:
                # Summary starts from the next line or two lines after
                extracted_summary = "\n".join(lines[summary_index + 1:])

        # Clean extracted summary from unnecessary blank lines
        extracted_summary = "\n".join([line for line in extracted_summary.splitlines() if line.strip()])

        # Create the custom resume DOCX file
        self.create_custom_resume_docx(template_path, custom_resume_path, extracted_job_title, extracted_summary, job_location)

        if create_pdf:
            # New code to convert DOCX to PDF
            custom_resume_pdf_path = custom_resume_path.replace('.docx', '.pdf')
            self.convert_docx_to_pdf(custom_resume_path, custom_resume_pdf_path)

        print(f"Custom resume saved to {custom_resume_path}")

    def coverletter_callback_function(self,output: TaskOutput):
        global job_details  # Ensure we're using the global dictionary

        print(self.job_details.get('company_name','unknown'))
        print(self.job_details.get('job_title','unknown'))

        # Extract the raw output from the task
        content = output.raw_output

    # Retrieve company name and job title from the global dictionary, replace with "unknown" if missing
        company_name = self.job_details.get('company_name', 'unknown')
        job_title = self.job_details.get('job_title', 'unknown')

        # Construct the nested folder path using company name and job title
        nested_folder_path = os.path.join(os.getcwd(), "jobs",company_name, job_title)

        # Check if the nested folder path exists
        if not os.path.exists(nested_folder_path):
            print(f"Error: The folder path {nested_folder_path} does not exist.")
            return

        # Construct the filename
        filename = f"Cover Letter - {candidate_name} - {company_name} - {job_title}.docx"
        # Define the file path for the cover letter inside the nested folder
        file_path = os.path.join(nested_folder_path, filename)


        # Create a new Word document
        doc = Document()
        
        # Modify the cover letter content to start with "Dear Hiring Manager,"
        # and remove any content before that if it exists
        content = output.raw_output
        start_index = content.find("Dear Hiring Manager,")
        if start_index != -1:
            content = content[start_index:]
        else:
            content = "Dear Hiring Manager,\n" + content
        
        # Split the content into lines
        lines = content.splitlines()
        
        # Check if the last line ends with ,,,
        if lines[-1].endswith("```"):
            # Remove the ,,, from the last line
            lines[-1] = lines[-1][:-3]
        
        # Reconstruct the content without ,,, on the last line
        content = "\n".join(lines)

        # Add the modified cover letter content to the document
        doc.add_paragraph(content)
        
        # Save the document
        doc.save(file_path)

        if create_pdf:
            # New code to convert DOCX to PDF
            cover_letter_pdf_path = file_path.replace('.docx', '.pdf')
            self.convert_docx_to_pdf(file_path, cover_letter_pdf_path)

        # Define the file path for the job description summary inside the nested folder
        file_path = os.path.join(nested_folder_path, "cover_letter.txt")

        # Write the content to the file
        try:
            with open(file_path, "w") as file:
                file.write(content)
            print(f"Cover letter saved to {file_path}")
        except IOError as e:
            print(f"Failed to write file {file_path}: {e}")


    def setup_agents_and_tasks(self):
      
        job_description_expert = Agent(
            role='Analyze  Job Description',
            goal='Analyze the job description to understand the company industry, role seniority level, main responsibilities, and identify the key skills required.',
            backstory="""You are expert in analyzing job descriptions and identifying the key skills required for the job""",
            verbose=True,
            llm=self.llm,
            allow_delegation=True,
            # tools=[description_read_tool]
            tools=[self.read_description_tool],
            memory=True
        )

        resume_expert = Agent(
            role='Analyze  Candidates\'s Resume',
            goal='Analyze a candidate\'s resume and highlight the key skills and achievements',
            backstory="""You are expert in analyzing resumes and identifying the key skills and achievements""",
            verbose=True,
            llm=self.llm,
            allow_delegation=True,
            tools=[self.resume_read_tool],
             memory=True
        )

        resume_writer = Agent(
            role='Resume and cover letter Writer',
            goal='Following best practices, write best resume and cover letter for the job description that matches the candidate\'s skill set, and yet tailored to the job description and the specific keywords used in the job description',
            backstory="""You are a resume and cover letter expert and can write the best resume and cover letter for the job description""",
            verbose=True,
            llm=self.llm,
            tools=[self.resume_read_tool],
            allow_delegation=True,
            memory=True
        )

        # Define tasks for your agents
        job_description_task = Task(
            description=f"""review the job description in path {self.output_file} with index={self.job_index}, using the description_read_tool tool. Analyze it and summarize the key responsibilities and skills they are looking for in the candidate. Follow the format: Company Name: <company>\n\n Job Location: <location>\n\n
            Job Title: <title>\n\n Responsibilities: <requirements>\n\n Skills: <skills>\n\n Company Industry: <industry>\n\n Role Seniority: <seniority>\n\n Role Discipline: <discipline>\n\n Salary Range: <range>""",
            agent=job_description_expert,
            # tools=[description_read_tool]
            tools=[self.read_description_tool],       
            callback=self.job_callback_function,
            # output_file="job_d
        )

        candidate_skills_task = Task(
            description=f"""Review {candidate_name}'s resume. You must use the resume_read_tool tool to review the resume. Analyze and highlight the skills and job experience he has. If there are skills achievements or titles in the resume that match what the job description is looking for, highlight them. Otherwise, choose from the resume the top background to highlight. Do not make up experience that is not either explicitly or strongly implied from the resume. Follow the format: Candidate's name: <name>\n\n Relevant Title: <title>\n\n Relevant skills: <skills>\n\n Relevant experience: <experience>""",
            agent=resume_expert,
            context=[job_description_task],
            tools=[self.resume_read_tool],
            # output_file="candidate_skills_summary.txt",
            callback=self.candidate_callback_function
        )

        resume_writing_task = Task(
            description=f"""Write a resume summary section that highlights {candidate_name}'s hard and soft skills. Try to highlight skills and keywords that match the job description. Do not make up skills or summary that do not truthfully represent {candidate_name}'s strength. The summary paragraph is about 5 sentences. Do not be too verbose, and make the sentences impactful, and do not use cliche words. When using sentences from {candidate_name}'s resume in the summary, do not copy word for word. But when applicable, it is ok to use exact keywords from the job description to show that the candidate is a great fit. Also Select a job title that closely matches the job title in the job description. The desired job title is 1-5 words. You must follow the format: Job Title: <title>\n\n Summary: <Summary>""",
            agent=resume_writer,
            
            context=[job_description_task, candidate_skills_task],
            # output_file="resume_summary.txt",
            callback=self.resume_callback_function
        )

        coverletter_writing_task = Task(
            description=f""" Write a cover letter that highlights {candidate_name}'s hard and soft skills. When a match found, highlight the areas that fit the skills and responsibilities required by the job description. To personalize the cover letter, make sure to mention the job title required in the job description and the company name if known, somewhere in the text. Do not make up skills or achievements that do not truthfully represent {candidate_name}'s strength. Structure the cover letter into 4 paragraphs: 1. Introduction, 2. Why you are interested in the job, 3. Why you are the best candidate for the job, 4. Conclusion. You must follow the format: <cover letter>. Start with Dear Hiring Manager, and end with Sincerely,\n {candidate_name}""",
            agent=resume_writer,
            context=[job_description_task, candidate_skills_task],
            # output_file="cover_letter.txt",
            callback=self.coverletter_callback_function
        )

        # Set up your crew with a sequential process (tasks executed sequentially by default)
        job_search_crew = Crew(
            agents=[job_description_expert, resume_expert,resume_writer],
            tasks=[ job_description_task, candidate_skills_task, resume_writing_task, coverletter_writing_task],
            process=Process.sequential,
            manager_llm=self.llm,
        )
        return job_search_crew


    def kickoff(self):
        crew_result =self.setup_agents_and_tasks().kickoff()
        print(crew_result)

        return self.job_details

if __name__ == "__main__":
    job_index= 0  # Example list of indices
    manager = JobSearchCrewManager(job_index)
    manager.kickoff()